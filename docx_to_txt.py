from docx import Document


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = []

    # 문서의 모든 문단을 순회하며 텍스트를 추출
    for para in doc.paragraphs:
        full_text.append(para.text)

    # 문서의 모든 표를 순회하며 텍스트를 추출
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            full_text.append('\t'.join(row_text))

    # 문서의 헤더와 푸터에서 텍스트 추출
        for section in doc.sections:
            header = section.header
            footer = section.footer
            for para in header.paragraphs:
                full_text.append(para.text)
            for para in footer.paragraphs:
                full_text.append(para.text)

    return '\n'.join(full_text)


# 사용 예시
# 여기에 .docx 파일 경로를 입력하세요.
file_path = r"D:\Users\ie-woo\Documents\Google 드라이브\docs\인터비즈시스템N\_작업\2022 0516a 다국어 번역사\@Translators-Pool-Search\abba\@test\test_sub\test_subsub\업데이트_이력서_박소운_20210113.docx"
extracted_text = extract_text_from_docx(file_path)
print(extracted_text)
