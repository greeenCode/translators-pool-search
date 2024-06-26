# -*- coding: utf-8 -*-

## 소스 폴더 내 pdf, html, docx에서 텍스트를 추출하여 OpenAI api 호출, 분석하여 번역사 정보를 엑셀파일에 
## 저장하는 프로젝트
# - 다수 파일처리를 위한 
#   - Checkpoint 저장 및 재개 기능 추가
#       - 저장 시 파일명+처리시간
#   - 에러 처리 및 로깅 기능 추가 - 에러 발생한 파일 정보를 기록 후 계속 진행
# - 엑셀 A1셀에 생성일자 기록
# - batch 처리
#   - api 호출 수를 줄이기 위해 한 번 호출에 다수(batch_size)의 추출 텍스트 전달
# - checkpoint.pkl에 file_data 저장, 엑셀저장 시 임시 .xlsx 생성
#   - log_error 추가하여 로그남기고 프로세스 멈추지 않도록
# - 정보항목에 '나이' 추가 -x

import os
import pandas as pd
from datetime import datetime
import openai
import yaml
import fitz  # PyMuPDF
import json
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font
import re
import pickle
import shutil
from docx import Document

from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

# Configurations
source_folder = r'D:\Users\ie-woo\Documents\Google 드라이브\docs\인터비즈시스템N\_작업\2022 0516a 다국어 번역사\@Translators-Pool-Search\abba'
target_folder = r'D:\Users\ie-woo\Documents\Google 드라이브\docs\인터비즈시스템N\_작업\2022 0516a 다국어 번역사\@Translators-Pool-Search'
target_path = os.path.join(target_folder, 'profiles_data5.xlsx')
temp_target_path = os.path.join(target_folder, 'profiles_data_temp.xlsx')
checkpoint_path = os.path.join(target_folder, 'checkpoint.pkl')
log_path = os.path.join(target_folder, 'error_log.txt')

# Load API Key from credentials.yml
with open('config/credentials.yml', 'r') as file:
    credentials = yaml.safe_load(file)
api_key = credentials['openai']['api_key']
openai.api_key = api_key

# Helper functions
def extract_text_from_pdf(file_path):
    text = ""
    try:
        with fitz.open(file_path) as doc:
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text("text")
    except Exception as e:
        log_error(f"Error extracting text from PDF {file_path}: {e}")
    return text

def extract_text_from_html(file_path):
    text = ""
    try:
        # Selenium 옵션 설정
        chrome_options = Options()
        chrome_options.add_argument("--headless")  # 브라우저 창을 띄우지 않음
        chrome_driver_path = r'C:\Util\chromedriver-win64\chromedriver.exe'  # ChromeDriver 경로로 변경
        service = Service(chrome_driver_path)

        # 브라우저 열기
        driver = webdriver.Chrome(service=service, options=chrome_options)

        try:
            # HTML 파일 열기
            driver.get(f'file:///{os.path.abspath(file_path)}')
            
            # 잠시 대기하여 페이지가 완전히 로드되도록 함
            WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.TAG_NAME, 'body')))
            
            # 페이지의 전체 텍스트 추출
            text = driver.find_element(By.TAG_NAME, 'body').text
            
        except Exception as e:
            log_error(f"Error extracting text from HTML {file_path}: {e}")
                       
        finally:
            # 드라이버 종료
            driver.close()
            driver.quit()
            
    except Exception as e:
        log_error(f"Error initializing WebDriver: {e}")
    
    return text

def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = Document(file_path)
        full_text = []
        
        # 문서의 모든 문단을 순회하며 텍스트를 추출
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # 문서의 모든 표를 순회하며 텍스트를 추출
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append('\t'.join(row_text))
        
        text = '\n'.join(full_text)
    except Exception as e:
        log_error(f"Error extracting text from DOCX {file_path}: {e}")
    return text

def format_multiline_text(text):
    if isinstance(text, list):
        text = '; '.join(text)
    return text.replace('; ', ';\n')

def clean_response_text(text):
    text = re.sub(r'[\x00-\x1F\x7F]', '', text)
    return text

def batch_extract_information(texts):
    batch_prompt_text = f"""
    주어진 텍스트를 분석해서 다음 정보 항목을 JSON 형식으로 추출해줘. 
    번역사의 이름, 이메일, 전화번호, 현 거주지(도시 이름까지만), 나이(출생년도가 표시되어 있으면 현재 년도까지 추정된 나이와 출생년도를 표기하고, 출생년도가 없으면 명시된 나이를 표기하고, 알수없으면 '알 수 없음'), 자기 소개 개요(공백 포함 300자 이내로), 번역 경력년수(명시되어있다면 명시된 년수와 활동 내역으로 추정한 시작년도를 표기하고, 명시되지 않았다면 활동 내역으로 추정한 시작년도부터 현재까지의 경과년수와 시작년도를 표기해줘. 알 수 없다면 '알 수 없음'), 번역 가능한 언어, 통역 가능한 언어, 번역 툴  사용가능여부(Trados, MemoQ, Smartcat 등 번역 툴 사용가능하면 툴 이름을 표기하고, 알수없으면 "알 수 없음"), 주요 학력, 주요 경력, 해외(한국 외) 교육기관에서 공부경험 유무(알 수 없다면 '알 수 없음'), 그밖에 번역사로서 경쟁력 등을 아래의 출력문 사례처럼 작성해줘.

    {{
        "이름": "양중남",
        "이메일": "inpraiseofdissent@gmail.com",
        "전화번호": "010-8751-1205",
        "거주지": "제주시",
        "나이": " 44세, 1980생"
        "자기소개": "양중남은 한영 통역 대학원을 졸업하고 미국에서 심리학 박사 학위를 받은 후 20년간 연구원으로 재직하였으며, 현재는 8년 경력의 프리랜서 번역사로 활동하고 있습니다. 영어와 한국어를 자유롭게 구사하는 바이링구얼 번역사입니다.",
        "경력년수": "8년 from 2001",
        "번역가능언어": "영어>한국어, 한국어>영어",
        "통역가능언어": "영어, 일본어",
        "번역툴가능여부": Trados, MemoQ
        "주요학력": 
        "박사, 실험 심리학, New York University (1995-1999); 
        학사, 영어 교육과, 제주 대학교 (1977-1981)",
        "주요경력": 
        "프리랜서 번역사 (2012-현재); 제주 대학 교육학과 대학원 강사 (2016-2018); 
        연구원, NASA Ames 연구소 (2002-2004); 박사후 과정, University of Chicago (1999-2001)",
        "해외학업유무": "New York University, Ball State University, University of Chicago",
        "경쟁력": 
        "다양한 분야 번역 경험 (자연과학, 사회과학, 비즈니스, 금융, 의학, 컴퓨터 과학 및 IT); 
        영한 자막 번역 경험 (의학, 교육, 음악, 드라마 등 다양한 분야에서 약 700개 비디오 번역); 
        미국에서 다수의 연구 논문 발표"
    }}

    여러 텍스트를 처리해야 하므로, 각 텍스트를 분석한 결과는 개별적으로 JSON 배열로 반환해줘.
    """
    prompt_texts = [f"텍스트:\n{text}" for text in texts]
    prompt = f"{batch_prompt_text}\n\n" + "\n\n".join(prompt_texts)
    
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a data extraction assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=3000,
        temperature=0.5
    )
    response_text = response['choices'][0]['message']['content'].strip()
    
    # Clean the response text
    response_text = response_text.replace("```json", "").replace("```", "").strip()
    response_text = clean_response_text(response_text)

    try:
        extracted_infos = json.loads(response_text)
        for extracted_info in extracted_infos:
            if '주요학력' in extracted_info:
                extracted_info['주요학력'] = format_multiline_text(extracted_info['주요학력'])
            if '주요경력' in extracted_info:
                extracted_info['주요경력'] = format_multiline_text(extracted_info['주요경력'])
            if '경쟁력' in extracted_info:
                extracted_info['경쟁력'] = format_multiline_text(extracted_info['경쟁력'])
            if '해외학업유무' in extracted_info:
                extracted_info['해외학업유무'] = format_multiline_text(extracted_info['해외학업유무'])
        return extracted_infos
    except json.JSONDecodeError as e:
        log_error(f"Error parsing JSON: {e}\nResponse text: {response_text}")
        return []

def log_error(message):
    with open(log_path, 'a') as log_file:
        log_file.write(f"{datetime.now().isoformat()} - {message}\n")

# Checkpoint loading
if os.path.exists(checkpoint_path):
    with open(checkpoint_path, 'rb') as f:
        checkpoint_data = pickle.load(f)
        processed_files_data = checkpoint_data.get('processed_files', [])
        current_batch_data = checkpoint_data.get('current_batch', [])
else:
    processed_files_data = []
    current_batch_data = []

# Main script
file_data = []
system_files = ['desktop.ini', 'Thumbs.db']
excluded_char = '@'
default_file_date = datetime(2000, 1, 1)

use_default_date = input(f"작업대상 파일의 수정일을 {default_file_date.date()} 이후로 하시겠습니까? (yes/y, [enter] to use default): ").strip().lower()
if use_default_date in ('yes', 'y', ''):
    modified_file_date = default_file_date
else:
    date_input = input("새로운 수정일을 입력하세요 (YYYY-MM-DD): ").strip()
    try:
        modified_file_date = datetime.strptime(date_input, '%Y-%m-%d')
    except ValueError:
        print("Invalid date format. Please enter the date in YYYY-MM-DD format.")
        exit()

# 기존 파일 수 카운팅
file_count = 0
processed_file_count = 0
for root, dirs, files in os.walk(source_folder):
    dirs[:] = [d for d in dirs if excluded_char not in d]
    files = [file for file in files if file not in system_files and excluded_char not in file and (file.lower().endswith('.pdf') or file.lower().endswith('.html') or file.lower().endswith('.docx'))]
    for file in files:
        file_path = os.path.join(root, file)
        file_modified_time = datetime.fromtimestamp(os.path.getmtime(file_path))
        if file_modified_time >= modified_file_date:
            file_count += 1
            if file_path in [item['file_path'] for item in processed_files_data]:
                processed_file_count += 1

# 실제 처리할 파일 수 계산
actual_file_count = file_count - processed_file_count

# 출력 부분
print(f"지정한 날자 이후의 pdf, html, docx 파일 수: {file_count}")
print(f"이미 처리된 파일 수: {processed_file_count}")
print(f"처리 대상 파일 수: {actual_file_count}")
proceed = input("계속 진행하시겠습니까?(yes/y, [enter] to continue): ").strip().lower()

if proceed in ('yes', 'y', ''):
    processed_count = 0
    batch_size = 5  # 배치 처리 크기 설정
    text_batch = []
    file_batch = []

    for root, dirs, files in os.walk(source_folder):
        dirs[:] = [d for d in dirs if excluded_char not in d]
        files = [file for file in files if file not in system_files and excluded_char not in file and (file.lower().endswith('.pdf') or file.lower().endswith('.html') or file.lower().endswith('.docx'))]
        for file in files:
            file_path = os.path.join(root, file)
            if file_path in [item['file_path'] for item in processed_files_data]:
                continue  # Skip already processed files

            file_modified_time = datetime.fromtimestamp(os.path.getmtime(file_path))
            if file_modified_time >= modified_file_date:
                processed_count += 1
                print(f"{processed_count}/{actual_file_count} 번째 파일 작업 중 ... (파일명: {file})")

                try:
                    if file.lower().endswith('.pdf'):
                        text = extract_text_from_pdf(file_path)
                    elif file.lower().endswith('.html'):
                        text = extract_text_from_html(file_path)
                    elif file.lower().endswith('.docx'):
                        text = extract_text_from_docx(file_path)
                    else:
                        continue

                    text_batch.append(text)
                    file_batch.append({
                        "file_path": file_path,
                        "file_modified_time": file_modified_time,
                        "file_name": file
                    })

                    if len(text_batch) >= batch_size:
                        extracted_infos = batch_extract_information(text_batch)
                        for idx, extracted_info in enumerate(extracted_infos):
                            if extracted_info is not None:
                                extracted_info['파일수정일'] = file_batch[idx]['file_modified_time'].strftime('%Y-%m-%d')
                                relative_file_path = os.path.relpath(file_batch[idx]['file_path'], start=target_folder)
                                extracted_info['File Link'] = f'=HYPERLINK("{relative_file_path}")'
                                file_data.append(extracted_info)

                                # Add to processed files data and save checkpoint
                                processed_files_data.append({'file_path': file_batch[idx]['file_path'], 'processed_time': datetime.now()})
                        with open(checkpoint_path, 'wb') as f:
                            pickle.dump({'processed_files': processed_files_data, 'current_batch': file_data}, f)
                        
                        # Reset batches
                        text_batch = []
                        file_batch = []

                except Exception as e:
                    log_error(f"Error processing file {file_path}: {e}")

    # 남은 배치 처리
    if text_batch:
        extracted_infos = batch_extract_information(text_batch)
        for idx, extracted_info in enumerate(extracted_infos):
            if extracted_info is not None:
                extracted_info['파일수정일'] = file_batch[idx]['file_modified_time'].strftime('%Y-%m-%d')
                relative_file_path = os.path.relpath(file_batch[idx]['file_path'], start=target_folder)
                extracted_info['File Link'] = f'=HYPERLINK("{relative_file_path}")'
                file_data.append(extracted_info)

                # Add to processed files data and save checkpoint
                processed_files_data.append({'file_path': file_batch[idx]['file_path'], 'processed_time': datetime.now()})

        with open(checkpoint_path, 'wb') as f:
            pickle.dump({'processed_files': processed_files_data, 'current_batch': file_data}, f)

    if file_data:
        try:
            if os.path.exists(target_path):
                wb = load_workbook(target_path)
                ws = wb.active
                # 기존 엑셀 파일의 1행 갱신
                current_time_text = datetime.now().strftime("On %Y-%m-%d %H:%M, the list below was last updated by OpenAI API")
                ws['A1'] = current_time_text
                blue_font = Font(color="0000FF")
                ws['A1'].font = blue_font
            else:
                wb = Workbook()
                ws = wb.active
                headers = ["이름", "File Link", "이메일", "전화번호", "거주지", "자기소개", "경력년수", "번역가능언어", "통역가능언어", "번역툴가능여부", "주요학력", "주요경력", "해외학업유무", "경쟁력", "파일수정일"]
                current_time_text = datetime.now().strftime("On %Y-%m-%d %H:%M, the list below was last updated by OpenAI API")
                ws.append([current_time_text])  # 1행에 생성일자 문구 추가
                ws.append(headers)  # 2행에 헤더를 기록
                blue_font = Font(color="0000FF")
                ws['A1'].font = blue_font

            for info in file_data:
                try:
                    row = [
                        info.get("이름", ""),
                        info.get("File Link", ""),
                        info.get("이메일", ""),
                        info.get("전화번호", ""),
                        info.get("거주지", ""),
                        info.get("자기소개", ""),
                        info.get("경력년수", ""),
                        info.get("번역가능언어", ""),
                        info.get("통역가능언어", ""),
                        format_multiline_text(info.get("번역툴가능여부", "")),
                        format_multiline_text(info.get("주요학력", "")),
                        format_multiline_text(info.get("주요경력", "")),
                        format_multiline_text(info.get("해외학업유무", "")),
                        format_multiline_text(info.get("경쟁력", "")),
                        info.get("파일수정일", "")
                    ]
                    ws.append(row)
                except Exception as e:
                    log_error(f"Error appending row for file {info.get('File Link', '')}: {e}")

            os.makedirs(target_folder, exist_ok=True)
            try:
                # 임시 파일에 먼저 저장
                wb.save(temp_target_path)
                
                # 임시 파일을 최종 파일로 이동
                shutil.move(temp_target_path, target_path)
                print(f"Finished! File information saved to '{target_path}'")
                
                # Save checkpoint without current batch
                with open(checkpoint_path, 'wb') as f:
                    pickle.dump({'processed_files': processed_files_data, 'current_batch': []}, f)
            except Exception as e:
                log_error(f"Error saving Excel file: {e}")
                if os.path.exists(temp_target_path):
                    os.remove(temp_target_path)
        except Exception as e:
            log_error(f"Error processing Excel file: {e}")
    else:
        print("No valid files processed.")
else:
    print("Process aborted by the user.")
