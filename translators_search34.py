# -*- coding: utf-8 -*-

## 소스 폴더 내 pdf, html, docx에서 텍스트를 추출하여 OpenAI api 호출, 분석하여 번역사 정보를 엑셀파일에 저장하는 프로젝트
# - 다수 파일처리에 에러발생 대비
#   - 에러 처리 및 로깅 기능 추가 - 에러 발생한 파일 정보를 기록 후 계속 진행
#   - 재시작 시 Checkpoint 사용않고 이전 결과물인 엑셀파일목록과 비교, 새 파일만 작업
# - batch 처리
# - selenium webdriver instance 재사용
# 오류 발견 - 빈 텍스트 추출 시, 에러 출력 후 다름 파일로 진행 
# 경과시간, 처리파일 수 출력

# -*- coding: utf-8 -*-

import os
import pandas as pd
from datetime import datetime
import openai
import yaml
import fitz  # PyMuPDF
import json
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font
import re
from docx import Document
import time  # 경과시간 측정을 위한 모듈

from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

# Configurations
source_folder = r'D:\Users\ie-woo\Documents\Google 드라이브\docs\인터비즈시스템N\_작업\2022 0516a 다국어 번역사\@Translators-Pool-Search\abba'
target_folder = r'D:\Users\ie-woo\Documents\Google 드라이브\docs\인터비즈시스템N\_작업\2022 0516a 다국어 번역사\@Translators-Pool-Search'
target_path = os.path.join(target_folder, 'profiles_data7.xlsx')
log_path = os.path.join(target_folder, 'error_log.txt')

# Load API Key from credentials.yml
with open('config/credentials.yml', 'r') as file:
    credentials = yaml.safe_load(file)
api_key = credentials['openai']['api_key']
openai.api_key = api_key

# Helper functions
def extract_text_from_pdf(file_path):
    text = ""
    try:
        with fitz.open(file_path) as doc:
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text("text")
        if not text.strip():  # 텍스트가 비어있는지 확인
            raise ValueError("Extracted text is empty")
    except Exception as e:
        log_error(f"Error extracting text from PDF {file_path}: {e}")
        print(f"Error extracting text from PDF {file_path}: {e}")
    return text

# WebDriver 초기화
chrome_options = Options()
chrome_options.add_argument("--headless")  # 브라우저 창을 띄우지 않음
chrome_driver_path = r'C:\Util\chromedriver-win64\chromedriver.exe'  # ChromeDriver 경로로 변경
service = Service(chrome_driver_path)
driver = webdriver.Chrome(service=service, options=chrome_options)

def extract_text_from_html(file_path, driver):
    text = ""
    try:
        try:
            # HTML 파일 열기
            driver.get(f'file:///{os.path.abspath(file_path)}')

            # 잠시 대기하여 페이지가 완전히 로드되도록 함
            WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.TAG_NAME, 'body')))

            # 페이지의 전체 텍스트 추출
            text = driver.find_element(By.TAG_NAME, 'body').text

            if not text.strip():  # 텍스트가 비어있는지 확인
                raise ValueError("Extracted text is empty")
        except Exception as e:
            log_error(f"Error extracting text from HTML {file_path}: {e}")
            print(f"Error extracting text from HTML {file_path}: {e}")

    except Exception as e:
        log_error(f"Error initializing WebDriver: {e}")
        print(f"Error initializing WebDriver: {e}")

    return text

def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = Document(file_path)
        full_text = []
        
        # 문서의 모든 문단을 순회하며 텍스트를 추출
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # 문서의 모든 표를 순회하며 텍스트를 추출
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append('\t'.join(row_text))
        
        text = '\n'.join(full_text)

        if not text.strip():  # 텍스트가 비어있는지 확인
            raise ValueError("Extracted text is empty")
    except Exception as e:
        log_error(f"Error extracting text from DOCX {file_path}: {e}")
        print(f"Error extracting text from DOCX {file_path}: {e}")
    return text

def format_multiline_text(text):
    if isinstance(text, list):
        text = '; '.join(text)
    return text.replace('; ', ';\n')

def clean_response_text(text):
    text = re.sub(r'[\x00-\x1F\x7F]', '', text)
    return text

def batch_extract_information(texts):
    batch_prompt_text = f"""
    주어진 텍스트를 분석해서 다음 정보 항목을 JSON 형식으로 추출해줘. 
    번역사의 이름, 이메일, 전화번호(+821027097063 > 010-2709-7063로 변환), 현 거주지(도시 이름까지만), 나이(출생년도가 표시되어 있으면 현재 년도까지 추정된 나이와 출생년도를 표기하고, 출생년도가 없으면 명시된 나이를 표기하고, 알수없으면 '알 수 없음'), 자기 소개 개요(공백 포함 400자 이내로), 번역 경력년수(명시되어있다면 명시된 년수와 활동 내역으로 추정한 시작년도를 표기하고, 명시되지 않았다면 활동 내역으로 추정한 시작년도부터 현재까지의 경과년수와 시작년도를 표기해줘. 알 수 없다면 '알 수 없음'), 번역 가능한 언어, 통역 가능한 언어, 번역 툴  사용가능여부(Trados, MemoQ, Smartcat 등 번역 툴 사용가능하면 툴 이름을 표기하고, 알수없으면 "알 수 없음"), 주요 학력, 주요 경력, 해외(한국 외) 교육기관에서 공부경험 유무(알 수 없다면 '알 수 없음'), 그밖에 번역사로서 경쟁력 등을 아래의 출력문 사례처럼 작성해줘.

    {{
        "이름": "양중남",
        "이메일": "inpraiseofdissent@gmail.com",
        "전화번호": "010-8751-1205",
        "거주지": "제주시",
        "나이": "44세, 1980년생",
        "자기소개": "양중남은 한영 통역 대학원을 졸업하고 미국에서 심리학 박사 학위를 받은 후 20년간 연구원으로 재직하였으며, 현재는 8년 경력의 프리랜서 번역사로 활동하고 있습니다. 영어와 한국어를 자유롭게 구사하는 바이링구얼 번역사입니다.",
        "경력년수": "8년 from 2001",
        "번역가능언어": "영어>한국어, 한국어>영어",
        "통역가능언어": "영어, 일본어",
        "번역툴가능여부": Trados, MemoQ
        "주요학력": 
        "박사, 실험 심리학, New York University (1995-1999); 
        학사, 영어 교육과, 제주 대학교 (1977-1981)",
        "주요경력": 
        "프리랜서 번역사 (2012-현재); 제주 대학 교육학과 대학원 강사 (2016-2018); 
        연구원, NASA Ames 연구소 (2002-2004); 박사후 과정, University of Chicago (1999-2001)",
        "해외학업유무": "New York University, Ball State University, University of Chicago",
        "경쟁력": 
        "다양한 분야 번역 경험 (자연과학, 사회과학, 비즈니스, 금융, 의학, 컴퓨터 과학 및 IT); 
        영한 자막 번역 경험 (의학, 교육, 음악, 드라마 등 다양한 분야에서 약 700개 비디오 번역); 
        미국에서 다수의 연구 논문 발표"
    }}

    여러 텍스트를 처리해야 하므로, 각 텍스트를 분석한 결과는 개별적으로 JSON 배열로 반환해줘.
    """
    prompt_texts = [f"텍스트:\n{text}" for text in texts]
    prompt = f"{batch_prompt_text}\n\n" + "\n\n".join(prompt_texts)
    
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a data extraction assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=3000,
        temperature=0.5
    )
    response_text = response['choices'][0]['message']['content'].strip()
    
    # Clean the response text
    response_text = response_text.replace("```json", "").replace("```", "").strip()
    response_text = clean_response_text(response_text)

    try:
        extracted_infos = json.loads(response_text)
        for extracted_info in extracted_infos:
            if '주요학력' in extracted_info:
                extracted_info['주요학력'] = format_multiline_text(extracted_info['주요학력'])
            if '주요경력' in extracted_info:
                extracted_info['주요경력'] = format_multiline_text(extracted_info['주요경력'])
            if '경쟁력' in extracted_info:
                extracted_info['경쟁력'] = format_multiline_text(extracted_info['경쟁력'])
            if '해외학업유무' in extracted_info:
                extracted_info['해외학업유무'] = format_multiline_text(extracted_info['해외학업유무'])
        return extracted_infos
    except json.JSONDecodeError as e:
        log_error(f"Error parsing JSON: {e}\nResponse text: {response_text}")
        return []

def log_error(message):
    with open(log_path, 'a') as log_file:
        log_file.write(f"{datetime.now().isoformat()} - {message}\n")

# Main script
file_data = []
system_files = ['desktop.ini', 'Thumbs.db']
excluded_char = '@'
default_file_date = datetime(2000, 1, 1)

use_default_date = input(f"작업대상 파일의 수정일을 {default_file_date.date()} 이후로 하시겠습니까? (yes/y, [enter] to use default): ").strip().lower()
if use_default_date in ('yes', 'y', ''):
    modified_file_date = default_file_date
else:
    date_input = input("새로운 수정일을 입력하세요 (YYYY-MM-DD): ").strip()
    try:
        modified_file_date = datetime.strptime(date_input, '%Y-%m-%d')
    except ValueError:
        print("Invalid date format. Please enter the date in YYYY-MM-DD format.")
        exit()

# 기존 파일 수 카운팅 및 처리 대상 파일 목록 생성
file_count = 0
file_to_process_list = []
for root, dirs, files in os.walk(source_folder):
    dirs[:] = [d for d in dirs if excluded_char not in d]
    files = [file for file in files if file not in system_files and excluded_char not in file and (file.lower().endswith('.pdf') or file.lower().endswith('.html') or file.lower().endswith('.docx'))]
    for file in files:
        file_path = os.path.join(root, file)
        file_modified_time = datetime.fromtimestamp(os.path.getmtime(file_path))
        if file_modified_time >= modified_file_date:
            file_count += 1
            file_to_process_list.append(file_path)

# 이미 처리된 파일 목록 생성
processed_file_list = set()
if os.path.exists(target_path):
    wb = load_workbook(target_path)
    ws = wb.active
    for row in ws.iter_rows(min_row=3, values_only=True):
        link = row[1]
        if link and link.startswith('=HYPERLINK'):
            # 파일 경로 추출
            file_path = link.split('"')[1]  # =HYPERLINK("path")에서 path 추출
            processed_file_list.add(file_path)


# 상대 경로로 변환된 리스트 생성
relative_file_to_process_list = [os.path.relpath(file, start=target_folder) for file in file_to_process_list]

# 실제 처리할 파일 목록 및 수 계산
actual_file_to_process_list = [file for file in file_to_process_list if os.path.relpath(file, start=target_folder) not in processed_file_list]

actual_file_count = len(actual_file_to_process_list)

# # 디버그 출력
# print("file_to_process_list:")
# for file in file_to_process_list:
#     print(file)

# print("\nrelative_file_to_process_list:")
# for file in relative_file_to_process_list:
#     print(file)

# print("\nprocessed_file_list:")
# for file in processed_file_list:
#     print(file)

# print("\nactual_file_to_process_list:")
# for file in actual_file_to_process_list:
#     print(file)

# 출력 부분
print(f"지정한 날자 이후의 pdf, html, docx 파일 수: {file_count}")
print(f"이미 처리된 파일 수: {file_count - actual_file_count}")
print(f"처리 대상 파일 수: {actual_file_count}")
# 처리할 파일 수가 0인 경우 종료
if actual_file_count == 0:
    print("처리할 파일이 없습니다!")
    exit()
proceed = input("계속 진행하시겠습니까?(yes/y, [enter] to continue): ").strip().lower()

# 기존 코드에서 변경된 부분

if proceed in ('yes', 'y', ''):
    start_time = time.time()  # 경과시간 측정을 위한 시작 시간 기록

    processed_count = 0
    batch_size = 5  # 배치 처리 크기 설정
    text_batch = []
    file_batch = []

    for file_path in actual_file_to_process_list:
        processed_count += 1
        print(f"{processed_count}/{actual_file_count} 번째 파일 작업 중 ... (파일명: {os.path.basename(file_path)})")

        try:
            if file_path.lower().endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif file_path.lower().endswith('.html'):
                text = extract_text_from_html(file_path, driver)  # driver 전달
            elif file_path.lower().endswith('.docx'):
                text = extract_text_from_docx(file_path)
            else:
                continue

            if text.strip():  # 텍스트가 비어있지 않으면 배치에 추가
                text_batch.append(text)
                file_batch.append({
                    "file_path": file_path,
                    "file_modified_time": datetime.fromtimestamp(os.path.getmtime(file_path)),
                    "file_name": os.path.basename(file_path)
                })

            if len(text_batch) >= batch_size:
                extracted_infos = batch_extract_information(text_batch)
                for idx, extracted_info in enumerate(extracted_infos):
                    if extracted_info is not None:
                        extracted_info['파일수정일'] = file_batch[idx]['file_modified_time'].strftime('%Y-%m-%d')
                        relative_file_path = os.path.relpath(file_batch[idx]['file_path'], start=target_folder)
                        extracted_info['File Link'] = f'=HYPERLINK("{relative_file_path}")'
                        file_data.append(extracted_info)

                        # Add to processed files set
                        processed_file_list.add(relative_file_path)
                
                # Save file_data to Excel
                if file_data:
                    try:
                        if os.path.exists(target_path):
                            wb = load_workbook(target_path)
                            ws = wb.active
                            # 기존 엑셀 파일의 1행 갱신
                            current_time_text = datetime.now().strftime("On %Y-%m-%d %H:%M, the list below was last updated by OpenAI API")
                            ws['A1'] = current_time_text
                            blue_font = Font(color="0000FF")
                            ws['A1'].font = blue_font
                        else:
                            wb = Workbook()
                            ws = wb.active
                            headers = ["이름", "File Link", "이메일", "전화번호", "거주지", "나이", "자기소개", "경력년수", "번역가능언어", "통역가능언어", "번역툴가능여부", "주요학력", "주요경력", "해외학업유무", "경쟁력", "파일수정일"]
                            current_time_text = datetime.now().strftime("On %Y-%m-%d %H:%M, the list below was last updated by OpenAI API")
                            ws.append([current_time_text])  # 1행에 생성일자 문구 추가
                            ws.append(headers)  # 2행에 헤더를 기록
                            blue_font = Font(color="0000FF")
                            ws['A1'].font = blue_font

                        for info in file_data:
                            try:
                                row = [
                                    info.get("이름", ""),
                                    info.get("File Link", ""),
                                    info.get("이메일", ""),
                                    info.get("전화번호", ""),
                                    info.get("거주지", ""),
                                    info.get("나이", ""),
                                    info.get("자기소개", ""),
                                    info.get("경력년수", ""),
                                    info.get("번역가능언어", ""),
                                    info.get("통역가능언어", ""),
                                    format_multiline_text(info.get("번역툴가능여부", "")),
                                    format_multiline_text(info.get("주요학력", "")),
                                    format_multiline_text(info.get("주요경력", "")),
                                    format_multiline_text(info.get("해외학업유무", "")),
                                    format_multiline_text(info.get("경쟁력", "")),
                                    info.get("파일수정일", "")
                                ]
                                ws.append(row)
                            except Exception as e:
                                log_error(f"Error appending row for file {info.get('File Link', '')}: {e}")

                        os.makedirs(target_folder, exist_ok=True)
                        wb.save(target_path)
                        
                        # Reset file_data after saving
                        file_data = []
                    except Exception as e:
                        log_error(f"Error processing Excel file: {e}")

                # Reset batches
                text_batch = []
                file_batch = []

        except Exception as e:
            log_error(f"Error processing file {file_path}: {e}")
            log_error(file_path)  # Add file path to error log

    # 남은 배치 처리
    if text_batch:
        extracted_infos = batch_extract_information(text_batch)
        for idx, extracted_info in enumerate(extracted_infos):
            if extracted_info is not None:
                extracted_info['파일수정일'] = file_batch[idx]['file_modified_time'].strftime('%Y-%m-%d')
                relative_file_path = os.path.relpath(file_batch[idx]['file_path'], start=target_folder)
                extracted_info['File Link'] = f'=HYPERLINK("{relative_file_path}")'
                file_data.append(extracted_info)

                # Add to processed files set
                processed_file_list.add(relative_file_path)

        if file_data:
            try:
                if os.path.exists(target_path):
                    wb = load_workbook(target_path)
                    ws = wb.active
                    current_time_text = datetime.now().strftime("On %Y-%m-%d %H:%M, the list below was last updated by OpenAI API")
                    ws['A1'] = current_time_text
                    blue_font = Font(color="0000FF")
                    ws['A1'].font = blue_font
                else:
                    wb = Workbook()
                    ws = wb.active
                    headers = ["이름", "File Link", "이메일", "전화번호", "거주지", "나이", "자기소개", "경력년수", "번역가능언어", "통역가능언어", "번역툴가능여부", "주요학력", "주요경력", "해외학업유무", "경쟁력", "파일수정일"]
                    current_time_text = datetime.now().strftime("On %Y-%m-%d %H:%M, the list below was last updated by OpenAI API")
                    ws.append([current_time_text])
                    ws.append(headers)
                    blue_font = Font(color="0000FF")
                    ws['A1'].font = blue_font

                for info in file_data:
                    try:
                        row = [
                            info.get("이름", ""),
                            info.get("File Link", ""),
                            info.get("이메일", ""),
                            info.get("전화번호", ""),
                            info.get("거주지", ""),
                            info.get("나이", ""),
                            info.get("자기소개", ""),
                            info.get("경력년수", ""),
                            info.get("번역가능언어", ""),
                            info.get("통역가능언어", ""),
                            format_multiline_text(info.get("번역툴가능여부", "")),
                            format_multiline_text(info.get("주요학력", "")),
                            format_multiline_text(info.get("주요경력", "")),
                            format_multiline_text(info.get("해외학업유무", "")),
                            format_multiline_text(info.get("경쟁력", "")),
                            info.get("파일수정일", "")
                        ]
                        ws.append(row)
                    except Exception as e:
                        log_error(f"Error appending row for file {info.get('File Link', '')}: {e}")

                os.makedirs(target_folder, exist_ok=True)
                wb.save(target_path)
                # print(f"Finished! File information saved to '{target_path}'")
            except Exception as e:
                log_error(f"Error processing Excel file: {e}")

    end_time = time.time()  # 경과시간 측정을 위한 종료 시간 기록
    processed_time = end_time - start_time  # 경과시간 계산
    processed_time_str = time.strftime('%H:%M:%S', time.gmtime(processed_time))

    # 저장된 파일 수 계산
    saved_files_count = 0
    if os.path.exists(target_path):
        wb = load_workbook(target_path)
        ws = wb.active
        saved_files_count = ws.max_row - 2  # 헤더와 첫 줄을 제외한 실제 데이터 줄 수 계산

    # 처리 결과 출력
    print(f"Finished! File information saved to '{target_path}'")  # 마지막에 한 번만 출력되도록 이동

    print(f"처리 대상 파일 수: {actual_file_count}")
    print(f"엑셀에 저장된 파일 수: {saved_files_count}")
    print(f"프로세스 경과시간: {processed_time_str}")

else:
    print("Process aborted by the user.")

# 드라이버 종료
driver.quit()
