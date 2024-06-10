from docx import Document
import os
import win32com.client

def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text += cell_text + "\n"
    return text

def extract_text_from_doc(doc_file):
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(doc_file)
    text = ""
    
    # 본문 텍스트 추출
    text += doc.Content.Text + "\n"
    
    # 표의 텍스트 추출
    for table in doc.Tables:
        for row in table.Rows:
            for cell in row.Cells:
                cell_text = cell.Range.Text.strip()
                if cell_text:
                    text += cell_text + "\n"
    
    doc.Close(False)
    word.Quit()
    return text

def save_text_to_file(text, output_file):
    with open(output_file, "w", encoding="utf-8") as file:
        file.write(text)

def main():
    # 처리할 파일 경로
    file_path = r"abba\장수진.docx"
    
    # 파일 확장자에 따라 처리
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        text = extract_text_from_docx(file_path)
    elif ext == ".doc":
        text = extract_text_from_doc(file_path)
    else:
        print(f"지원하지 않는 파일 형식입니다: {ext}")
        return
    
    # 저장할 파일명 설정
    output_file = os.path.splitext(file_path)[0] + ".txt"
    
    # 텍스트를 파일로 저장
    save_text_to_file(text, output_file)
    
    print(f"텍스트가 {output_file}로 성공적으로 추출되었습니다.")

if __name__ == "__main__":
    main()
