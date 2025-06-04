# -*- coding: utf-8 -*-

import os
import pandas as pd
from datetime import datetime
from openai import OpenAI
import yaml
import fitz  # PyMuPDF
import json
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font
import re
from docx import Document
import docx2txt
import time

from selenium import webdriver
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

# Configurations
source_folder = r"D:\Users\ie-woo\Documents\인터비즈시스템N\_작업\2022 0516a 다국어 번역사\@Translators-Pool-Search\abba\@test\test_sub\test_subsub"
target_folder = r"D:\Users\ie-woo\Documents\인터비즈시스템N\_작업\2022 0516a 다국어 번역사\@Translators-Pool-Search\abba\@test\test_sub\test_subsub"
target_path = os.path.join(target_folder, f'test_profiles.xlsx')
log_path = os.path.join(target_folder, 'error_log.txt')

total_processed_tokens = 0
total_processed_cost = 0.0

# Load API Key from credentials.yml
with open('config/credentials.yml', 'r') as file:
    credentials = yaml.safe_load(file)
api_key = credentials['openai']['api_key']
client = OpenAI(api_key=api_key)

# WebDriver 초기화
chrome_options = Options()
# chrome_options.add_argument("--headless")
service = Service(ChromeDriverManager().install())
driver = webdriver.Chrome(service=service, options=chrome_options)

# Helper functions


def extract_text_from_pdf(file_path):
    text = ""
    try:
        with fitz.open(file_path) as doc:
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text("text")
        if not text.strip():
            raise ValueError("Extracted text is empty")
    except Exception as e:
        log_error(f"Error extracting text from PDF {file_path}: {e}")
        return None
    return text


def extract_text_from_html(file_path, driver):
    text = ""
    try:
        driver.get(f'file:///{os.path.abspath(file_path)}')
        WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.TAG_NAME, 'body')))
        text = driver.find_element(By.TAG_NAME, 'body').text
        if not text.strip():
            raise ValueError("Extracted text is empty")
    except Exception as e:
        log_error(f"Error extracting text from HTML {file_path}: {e}")
        return None
    return text


def extract_text_from_docx(file_path):
    def extract_text_with_docx2txt(file_path):
        try:
            return docx2txt.process(file_path)
        except Exception as e:
            return ''

    doc = Document(file_path)
    full_text = [para.text for para in doc.paragraphs]

    for table in doc.tables:
        for row in table.rows:
            full_text.append('\t'.join([cell.text for cell in row.cells]))

    header_footer_text = []
    for section in doc.sections:
        for para in section.header.paragraphs:
            header_footer_text.append(para.text)
        for para in section.footer.paragraphs:
            header_footer_text.append(para.text)

    full_text.extend(header_footer_text)
    return '\n'.join(full_text) if full_text else extract_text_with_docx2txt(file_path)


def format_multiline_text(text):
    if isinstance(text, list):
        return ',\n '.join(text)
    elif isinstance(text, dict):
        return '\n'.join([f"{key}: {value}" for key, value in text.items()])
    return text.replace('; ', ';\n').replace('\n ', '\n')


def extract_information(text):
    start_time = time.time()

    prompt = f"""
    주어진 텍스트에서 번역사 정보를 분석하여 다음 항목을 JSON 형식으로 추출해줘. JSON 형식의 사례는 아래와 같다.

    번역사의 이름(이름은 번역하지 말고 사용한 언어 그대로 써줘), 이메일, 전화번호(+821027097063 > 010-2709-7063로 변환, 알수없으면 '알 수 없음'), 현 거주지(도시 이름까지만, 알수없으면 '알 수 없음'), 나이(출생년도가 표시되어 있으면 현재 년도까지 추정된 나이와 출생년도를 표기하고, 출생년도가 없으면 명시된 나이를 표기하고, 알 수없으면 '알 수 없음'), 자기 소개 개요(프로필 내용을 바탕으로 400자 이내로 가능한 자세히 요약), 번역 경력년수(명시되어있다면 명시된 년수와 활동 내역으로 추정한 시작년도를 표기하고, 명시되지 않았다면 활동 내역으로 추정한 시작년도부터 현재까지의 경과년수와 시작년도를 표기해줘. 알 수 없다면 '알 수 없음'), 번역 가능한 언어(한국어>영어 처럼 반드시 언어방향 명시, "불어"라는 단어는 "프랑스어"로 바꿔써), 통역 가능한 언어("불어"라는 단어는 "프랑스어"로 바꿔써), 번역 툴  사용가능여부(Trados, MemoQ, Smartcat 등 번역 툴 사용가능하면 툴 이름을 표기하고, 알수없으면 "알 수 없음"), 주요 학력, 주요 경력, 해외(한국 외) 교육기관에서 공부경험 유무(알 수 없다면 '알 수 없음'), 그밖에 번역사로서 경쟁력(주요학력과 주요경력, 그 밖의 정보를 바탕으로 400자 이내로 가능한 자세히 요약 ), 통번역분야(경력을 바탕으로 번역, 통역 가능한 분야의 키워드와 사례, 주요 발주 기업명 포함, 빈도순으로 최대 10개 분야까지)

    JSON 형식의 사례:
    {{
        "이름": "양중남",
        "이메일": "inpraiseofdissent@gmail.com",
        "전화번호": "010-8751-1205",
        "거주지": "제주시",
        "나이": "44세, 1980년생",
        "자기소개": "양중남은 한영 통역 대학원을 졸업하고 미국에서 심리학 박사 학위를 받은 후 20년간 연구원으로 재직하였으며, 현재는 8년 경력의 프리랜서 번역사로 활동하고 있습니다. 영어와 한국어를 자유롭게 구사하는 바이링구얼 번역사입니다.",
        "경력년수": "8년 from 2001",
        "번역가능언어": ["영어>한국어", "한국어>영어"],
        "통역가능언어": ["영어", "일본어"],
        "번역툴가능여부": ["Trados", "MemoQ"],
        "주요학력": [
            "박사, 실험 심리학, New York University (1995-1999)",
            "학사, 영어 교육과, 제주 대학교 (1977-1981)"
        ],
        "주요경력": [
            "프리랜서 번역사 (2012-현재)",
            "제주 대학 교육학과 대학원 강사 (2016-2018)",
            "연구원, NASA Ames 연구소 (2002-2004)"
        ],
        "해외학업유무": ["New York University", "Ball State University", "University of Chicago"],
        "경쟁력": "다양한 분야 번역 경험 (자연과학, 사회과학, 비즈니스, 금융, 의학, 컴퓨터 과학 및 IT); 영한 자막 번역 경험 (의학, 교육, 음악, 드라마 등 다양한 분야에서 약 700개 비디오 번역); 미국에서 다수의 연구 논문 발표",
        "통번역분야": [
            "IT - 교육, 세미나 통역 (SAP, IBM, LG CNS, Oracle)",
            "금융 - 계약서, 보고서 번역 (한국씨티은행, Morgan Stanley, Bank of America)",
            "자동차 - 매뉴얼, 사용자 설명서 번역 (Mercedes-Benz, BMW, MAN Truck & Bus, Volkswagen, Audi)",
            "게임 현지화 - 컴투스㈜ (서머너즈 워, 소울시커, 사커 스피리츠 등)"
        ],
    }}

    텍스트: {text}
    """

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a data extraction assistant."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=3500,
            temperature=0.5
        )
        response_text = response.choices[0].message.content.strip()
        response_text = response_text.replace(
            "```json", "").replace("```", "").strip()

    except Exception as e:
        log_error(f"OpenAI API error: {e}")
        print(f"OpenAI API error: {e}")
        print("Process terminated due to API error.")
        exit(1)

    try:
        extracted_info = json.loads(response_text)
        if '번역가능언어' in extracted_info:
            extracted_info['번역가능언어'] = format_multiline_text(
                extracted_info['번역가능언어'])
        if '통역가능언어' in extracted_info:
            extracted_info['통역가능언어'] = format_multiline_text(
                extracted_info['통역가능언어'])
        if '번역툴가능여부' in extracted_info:
            extracted_info['번역툴가능여부'] = format_multiline_text(
                extracted_info['번역툴가능여부'])
        if '주요학력' in extracted_info:
            extracted_info['주요학력'] = format_multiline_text(
                extracted_info['주요학력'])
        if '주요경력' in extracted_info:
            extracted_info['주요경력'] = format_multiline_text(
                extracted_info['주요경력'])
        if '경쟁력' in extracted_info:
            extracted_info['경쟁력'] = format_multiline_text(
                extracted_info['경쟁력'])
        if '해외학업유무' in extracted_info:
            extracted_info['해외학업유무'] = format_multiline_text(
                extracted_info['해외학업유무'])
        if '통번역분야' in extracted_info:
            extracted_info['통번역분야'] = format_multiline_text(
                extracted_info['통번역분야'])

        # Calculate and print the number of tokens used and cost
        prompt_tokens = response.usage.prompt_tokens
        completion_tokens = response.usage.completion_tokens
        total_tokens = response.usage.total_tokens

        input_cost = (prompt_tokens / 1_000_000) * 0.15
        output_cost = (completion_tokens / 1_000_000) * 0.6
        processed_cost = input_cost + output_cost

        global total_processed_tokens
        global total_processed_cost
        total_processed_tokens += total_tokens
        total_processed_cost += processed_cost

        end_time = time.time()
        processed_time = end_time - start_time
        processed_time_str = f"{
            int(processed_time // 60):02}:{int(processed_time % 60):02}"

        print(f"Prompt tokens: {prompt_tokens}")
        print(f"Completion tokens: {completion_tokens}")
        print(f"Total tokens: {total_tokens}")
        print(f"batch 프로세스 경과시간: {processed_time_str}")
        print(f"Cost: ${processed_cost:.5f}")
        print(f"\n")

        return extracted_info

    except json.JSONDecodeError as e:
        log_error(f"Error parsing JSON: {e}\nResponse text: {response_text}")
        return []


def log_error(message):
    with open(log_path, 'a', encoding='utf-8') as log_file:
        log_file.write(f"{datetime.now().isoformat()} - {message}\n")


def save_to_excel(file_data, target_path):
    try:
        if os.path.exists(target_path):
            wb = load_workbook(target_path)
            ws = wb.active
        else:
            wb = Workbook()
            ws = wb.active
            headers = ["이름", "File Link", "이메일", "전화번호", "거주지", "나이", "자기소개", "경력년수",
                       "번역가능언어", "통역가능언어", "번역툴가능여부", "주요학력", "주요경력", "해외학업유무", "경쟁력", "통번역분야", "파일수정일"]
            ws.append(headers)

        for info in file_data:
            try:
                row = [
                    info.get("이름", ""),
                    info.get("File Link", ""),
                    info.get("이메일", ""),
                    info.get("전화번호", ""),
                    info.get("거주지", ""),
                    info.get("나이", ""),
                    info.get("자기소개", ""),
                    info.get("경력년수", ""),
                    format_multiline_text(info.get("번역가능언어", "")),
                    format_multiline_text(info.get("통역가능언어", "")),
                    format_multiline_text(info.get("번역툴가능여부", "")),
                    format_multiline_text(info.get("주요학력", "")),
                    format_multiline_text(info.get("주요경력", "")),
                    format_multiline_text(info.get("해외학업유무", "")),
                    format_multiline_text(info.get("경쟁력", "")),
                    format_multiline_text(info.get("통번역분야", "")),
                    info.get("파일수정일", "")
                ]
                ws.append(row)
            except Exception as e:
                log_error(f"Error appending row for file {
                          info.get('File Link', '')}: {e}")

        os.makedirs(target_folder, exist_ok=True)
        wb.save(target_path)
    except Exception as e:
        log_error(f"Error processing Excel file: {e}")


# Main script
main_start_time = time.time()
file_data = []
system_files = ['desktop.ini', 'Thumbs.db']
excluded_char = '@'
default_file_date = datetime(2000, 1, 1)

use_default_date = input(
    f"작업대상 파일의 수정일을 {default_file_date.date()} 이후로 하시겠습니까? (yes/y, [enter] to use default): ").strip().lower()
if use_default_date in ('yes', 'y', ''):
    modified_file_date = default_file_date
else:
    date_input = input("새로운 수정일을 입력하세요 (YYYY-MM-DD): ").strip()
    try:
        modified_file_date = datetime.strptime(date_input, '%Y-%m-%d')
    except ValueError:
        print("Invalid date format. Please enter the date in YYYY-MM-DD format.")
        exit()

# 소스 폴더 파일 수 카운팅 및 처리 대상 파일 목록 생성
file_count = 0
file_to_process_list = []
for root, dirs, files in os.walk(source_folder):
    dirs[:] = [d for d in dirs if excluded_char not in d]
    files = [file for file in files if file not in system_files and excluded_char not in file and (
        file.lower().endswith('.pdf') or file.lower().endswith('.html') or file.lower().endswith('.docx'))]
    for file in files:
        file_path = os.path.join(root, file)
        file_modified_time = datetime.fromtimestamp(
            os.path.getmtime(file_path))
        if file_modified_time >= modified_file_date:
            file_count += 1
            file_to_process_list.append(file_path)

# 이미 처리된 파일 목록 생성
processed_file_list = set()
if os.path.exists(target_path):
    wb = load_workbook(target_path)
    ws = wb.active
    for row in ws.iter_rows(min_row=2, values_only=True):
        link = row[1]
        if link and link.startswith('=HYPERLINK'):
            file_path = link.split('"')[1]
            processed_file_list.add(file_path)

# 상대 경로로 변환된 리스트 생성
relative_file_to_process_list = [os.path.relpath(
    file, start=target_folder) for file in file_to_process_list]

# 실제 처리할 파일 목록 및 수 계산
actual_file_to_process_list = [file for file in file_to_process_list if os.path.relpath(
    file, start=target_folder) not in processed_file_list]

actual_file_count = len(actual_file_to_process_list)

# 출력 부분
print(f"지정일 이후 처리대상 파일 수: {file_count}")
print(f"이미 처리된 파일 수: {file_count - actual_file_count}")
print(f"실제 처리 대상 파일 수: {actual_file_count}")

# 처리할 파일 수가 0인 경우 종료
if actual_file_count == 0:
    print("처리할 파일이 없습니다!")
    exit()

# 계속 진행할지 묻기
proceed = input("계속 진행하시겠습니까?(yes/y, [enter] to continue): ").strip().lower()

if proceed in ('yes', 'y', ''):
    processed_count = 0

    for file_path in actual_file_to_process_list:
        processed_count += 1
        print(
            f"{processed_count}/{actual_file_count} 번째 파일 작업 중 ... (파일명: {os.path.basename(file_path)})")

        try:
            if file_path.lower().endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif file_path.lower().endswith('.html'):
                text = extract_text_from_html(file_path, driver)
            elif file_path.lower().endswith('.docx'):
                text = extract_text_from_docx(file_path)
            else:
                continue

            if text is None or not text.strip():
                continue

            if text:
                extracted_info = extract_information(text)
                if extracted_info:
                    file_modified_time = datetime.fromtimestamp(
                        os.path.getmtime(file_path))
                    extracted_info['파일수정일'] = file_modified_time.strftime(
                        '%Y-%m-%d')
                    relative_file_path = os.path.relpath(
                        file_path, start=target_folder)
                    extracted_info['File Link'] = f'=HYPERLINK("{
                        relative_file_path}")'
                    file_data.append(extracted_info)

        except Exception as e:
            log_error(f"Error processing file {file_path}: {e}")
            log_error(file_path)

else:
    print("Process aborted by the user.")

# Save file_data to Excel
if file_data:
    save_to_excel(file_data, target_path)
    file_data = []

# 저장된 파일 수 계산
saved_files_count = 0
if os.path.exists(target_path):
    wb = load_workbook(target_path)
    ws = wb.active
    saved_files_count = ws.max_row - 1

# main 함수 경과시간 계산
main_end_time = time.time()
main_processed_time = main_end_time - main_start_time
main_processed_time_str = f"{
    int(main_processed_time // 60):02}:{int(main_processed_time % 60):02}"

# 처리 결과 출력
print(f"Finished! File information saved to '{target_path}'")
print(f"지정일 이후 처리대상 총 파일 수: {file_count}")
print(f"실제 처리 대상 파일 수: {actual_file_count}")
print(f"엑셀에 저장된 총 파일 수: {saved_files_count}")
print(f"\n")
print(f"Total processed tokens: {total_processed_tokens}")
print(f"Total processed Cost: ${total_processed_cost:.5f}")
print(f"프로세스 경과시간: {main_processed_time_str}")

# WebDriver 종료
driver.quit()
