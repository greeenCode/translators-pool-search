# docx  > txt, docx라이브러리로 추출이 empty일 때 2차로 docx2txt 라이브러리 사용

from docx import Document
import docx2txt
import os


def extract_text_from_docx(file_path):
    def extract_text_with_docx2txt(file_path):
        try:
            return docx2txt.process(file_path)
        except Exception as e:
            print(f'Error using docx2txt: {e}')
            return ''

    def is_list_all_whitespace(lst):
        return all(not item.strip() for item in lst)

    doc = Document(file_path)
    full_text = []

    # 문서의 모든 문단을 순회하며 텍스트를 추출
    for para in doc.paragraphs:
        full_text.append(para.text)

    # 문서의 모든 표를 순회하며 텍스트를 추출
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            full_text.append('\t'.join(row_text))

    # 문서의 헤더와 푸터에서 텍스트 추출
    for section in doc.sections:
        header = section.header
        footer = section.footer
        for para in header.paragraphs:
            full_text.append(para.text)
        for para in footer.paragraphs:
            full_text.append(para.text)
            # full_text = None    # None 테스트

    if not full_text or is_list_all_whitespace(full_text):
        docx2txt_text = extract_text_with_docx2txt(file_path)
        if docx2txt_text:
            full_text = docx2txt_text.splitlines()
            # print(f'docx2txt 추출: {full_text}')

    return '\n'.join(full_text)


def save_text_to_file(text, file_path):
    text_file_path = os.path.splitext(file_path)[0] + '.txt'
    with open(text_file_path, 'w', encoding='utf-8') as text_file:
        text_file.write(text)


# 사용 예시
# 여기에 .docx 파일 경로를 입력하세요.
file_path = r"D:\Users\ie-woo\Documents\Google 드라이브\docs\인터비즈시스템N\_작업\2022 0516a 다국어 번역사\abba resource\2020 지원자\1.영어\김영주.docx"
extracted_text = extract_text_from_docx(file_path)
save_text_to_file(extracted_text, file_path)
# print(f'추출 텍스트: {extracted_text}')
print(f'추출된 텍스트가 저장되었습니다: {os.path.splitext(file_path)[0] + ".txt"}')
